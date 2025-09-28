from fastapi import FastAPI, HTTPException
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from datetime import date
import os, uuid

try:
    from docx import Document
except Exception:
    Document = None

try:
    from weasyprint import HTML
except Exception:
    HTML = None

app = FastAPI(title="Siddarth Digital Twin Backend")
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Deliverable(BaseModel):
    title: str
    type: str
    markdown: str
    docx: Optional[str] = ""
    pdf: Optional[str] = ""

class TwinResponse(BaseModel):
    intake_summary: str
    plan_markdown: str
    plan_gantt_mermaid: str
    mindmap_mermaid: str
    architecture_mermaid: str
    process_mermaid: str
    deliverables: List[Deliverable]
    raid_table_markdown: str
    decision_log_markdown: str
    export_suggestions: List[str]

class TwinRequest(BaseModel):
    userPrompt: str
    deliverableType: Optional[str] = None

class ExportDocxRequest(BaseModel):
    title: str
    markdown: str

class ExportPdfRequest(BaseModel):
    html: str

def LLM_STUB(user_prompt: str, deliverable_type: Optional[str]) -> TwinResponse:
    today = str(date.today())
    intake = "\n".join([
        "- Objective: Generate artefacts and visuals",
        f"- User request: {user_prompt}",
        f"- Deliverable bias: {deliverable_type or 'auto'}",
        "- Assumptions: scope and stakeholders to be confirmed",
        "- Success: actionable templates + plan + mindmap + diagrams",
    ])
    plan_md = (
        "**Phase 1 – Discovery:** Stakeholders, problem framing, constraints\n"
        "**Phase 2 – Options & Roadmap:** Options analysis, backlog, milestones\n"
        "**Phase 3 – Build-Measure:** Iterations with stories/AC, diagrams\n"
        "**Phase 4 – UAT & Readiness:** Test strategy, training, comms\n"
        "**Phase 5 – Go-Live & Hypercare:** Cutover plan, KPIs, monitoring"
    )
    gantt = f"""```mermaid\ngantt\n  title Delivery Plan (High Level)\n  dateFormat  YYYY-MM-DD\n  section Initiation\n  Discovery & Stakeholders        :active, a1, {today}, 7d\n  Business Case & Options         :a2, after a1, 7d\n  section Delivery Setup\n  Roadmap & Backlog               :a3, after a2, 10d\n  Architecture & Interfaces       :a4, after a2, 10d\n  section Build & Test\n  Iteration 1                     :a5, after a3, 14d\n  Iteration 2                     :a6, after a5, 14d\n  UAT & Readiness                 :a7, after a6, 10d\n  section Launch\n  Go-Live & Hypercare             :a8, after a7, 7d\n```"""
    mindmap = """```mermaid\nmindmap\n  root((Goal))\n    Context\n      Business drivers\n      Users & personas\n      Constraints\n    Workstreams\n      BA\n        Discovery\n        Requirements\n        BPMN\n      PM\n        Plan\n        RAID\n        Comms\n      Agile\n        Roadmap\n        Epics/Stories\n      Solutions\n        Options\n        Integration spec\n      ML\n        Data readiness\n        Baseline & metrics\n        MLOps\n    Deliverables\n      BRD\n      WBS\n      Architecture diagram\n      Test strategy\n      Change plan\n```"""
    arch = """```mermaid\nflowchart LR\n  U[End Users] -->|Web/App| FE[Frontend (Next.js)]\n  FE --> API[(Backend API - FastAPI)]\n  API --> INT[Integration Layer / iPaaS]\n  API --> DB[(Operational DB)]\n  INT --> SaaS[(SaaS/ERP/CRM)]\n  INT --> DWH[(Analytics DWH)]\n  API --> AUTH[(AuthN/Z)]\n  subgraph ML\n    FS[Feature Store] --> MLR[Model Registry]\n    MLR --> SVC[Model Serving]\n  end\n  DB --> FS\n  SVC --> API\n```"""
    proc = """```mermaid\nflowchart TD\n  A[Capture Requirement] --> B[Classify (BA/PM/Agile/SE/ML)]\n  B -->|BA| C[Generate BRD Template]\n  B -->|PM| D[Generate WBS/RAID]\n  B -->|Agile| E[Epics/Stories/AC]\n  B -->|SE| F[Options & Interface Spec]\n  B -->|ML| G[ML Canvas & Pipeline]\n  C & D & E & F & G --> H[Assemble Pack + Diagrams + Mindmap]\n  H --> I{User Review}\n  I -->|Refine| B\n  I -->|Approve| J[Export DOCX/PDF/MD]\n```"""
    deliverables = [
        Deliverable(
            title="Business Requirements Document (Draft)",
            type="BRD",
            markdown=(
                f"# BRD\n\n## 1. Executive Summary\nDrafted from:\n> {user_prompt}\n\n"
                "## 2. Scope\n- In-scope: ...\n- Out-of-scope: ...\n\n"
                "## 3. Stakeholders\n- Sponsor, SMEs, Users\n\n"
                "## 4. Functional Requirements\n- FR-001 ...\n\n"
                "## 5. Non-Functional Requirements\n- Availability, Security, Privacy, Accessibility\n\n"
                "## 6. Risks & Assumptions\n- R-01 ...\n\n"
                "## 7. Acceptance Criteria\n- Given/When/Then ...\n\n"
                "## 8. Next Steps\n- Clarify data sources, integrations, constraints\n"
            ),
        ),
        Deliverable(
            title="RAID Log (Initial)",
            type="RAID",
            markdown=(
                "| ID | Type | Description | Likelihood | Impact | Owner | Mitigation |\n"
                "|---|---|---|---|---|---|---|\n"
                "| R-01 | Risk | Data quality issues | Medium | High | BA | Profiling + validation |\n"
                "| A-01 | Assumption | Vendor APIs stable | n/a | n/a | SE | Confirm in RFI |\n"
                "| I-01 | Issue | Env parity gaps | Medium | Medium | Eng | Align IaC |\n"
                "| D-01 | Decision | SaaS over custom build | n/a | n/a | SteerCo | TCO vs time-to-value |"
            ),
        ),
    ]
    return TwinResponse(
        intake_summary=intake,
        plan_markdown=plan_md,
        plan_gantt_mermaid=gantt,
        mindmap_mermaid=mindmap,
        architecture_mermaid=arch,
        process_mermaid=proc,
        deliverables=deliverables,
        raid_table_markdown=deliverables[1].markdown,
        decision_log_markdown=(
            "| Date | Decision | Context | Alternatives | Rationale |\n|---|---|---|---|---|\n| YYYY-MM-DD | Adopt SaaS | Time-to-value | Custom, Hybrid | Lower TCO, faster delivery |"
        ),
        export_suggestions=["DOCX", "PDF", "Markdown"],
    )

@app.post("/api/generate", response_model=TwinResponse)
def generate(req: TwinRequest):
    return LLM_STUB(req.userPrompt, req.deliverableType)

@app.post("/api/export/docx")
def export_docx(payload: ExportDocxRequest):
    if Document is None:
        raise HTTPException(status_code=500, detail="python-docx not installed. Run: pip install python-docx")
    doc = Document()
    doc.add_heading(payload.title, level=1)
    for line in payload.markdown.split("\n"):
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            doc.add_paragraph(line)
    out_path = f"/tmp/{uuid.uuid4().hex}.docx"
    doc.save(out_path)
    return {"path": out_path, "message": "DOCX created", "filename": os.path.basename(out_path)}

@app.post("/api/export/pdf")
def export_pdf(payload: ExportPdfRequest):
    if HTML is None:
        raise HTTPException(status_code=500, detail="WeasyPrint not installed. Run: pip install weasyprint")
    out_path = f"/tmp/{uuid.uuid4().hex}.pdf"
    HTML(string=payload.html).write_pdf(out_path)
    return {"path": out_path, "message": "PDF created", "filename": os.path.basename(out_path)}

@app.get("/health")
def health():
    return {"ok": True}
